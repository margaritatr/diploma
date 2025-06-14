from flask import Flask, request, render_template, redirect, jsonify, url_for,flash, send_file, abort, session
from flask_cors import CORS
from flask_login import LoginManager, login_user, login_required, logout_user, current_user
from flask_bcrypt import Bcrypt
import datetime
from datetime import datetime, timedelta
import calendar
from flask import make_response
from docx import Document
from io import BytesIO
from DAL import ORM, User, Client, Service, ServiceType, FunctionalClass, TerritorialCenter, Department, FunctionalClassServiceType
from docx.shared import Pt
from functools import wraps
import os


app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = os.urandom(24)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

bcrypt = Bcrypt(app)
orm = ORM()
orm.start_session()

@login_manager.user_loader
def load_user(user_id):
    return orm.get_by_id(User, user_id)

def role_required(role):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if current_user.role.role_name != role:
                abort(403)
            return f(*args, **kwargs)
        return decorated_function
    return decorator

def nocache(view):
    @wraps(view)
    def no_cache(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for('login'))
        response = make_response(view(*args, **kwargs))
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response
    return no_cache


@app.route('/')
def home():
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = orm.get_data(User).filter_by(username=username).first()

        if user and bcrypt.check_password_hash(user.password, password):
            login_user(user)
            if user.role.role_name == 'inspector':
                return redirect(url_for('inspector_welcome'))
            return redirect(url_for('select_month'))
        else:
            return render_template('login.html', error='Неверное имя пользователя или пароль')

    return render_template('login.html')

@app.route('/logout')
@nocache
def logout():
    logout_user()
    session.clear()
    return redirect(url_for('login'))


# СОЦИАЛЬНЫЙ РАБОТНИК


def generate_available_months(start_year):
    current_year = datetime.now().year
    current_month = datetime.now().month
    available_months = {}

    for year in range(current_year, start_year - 1, -1):
        available_months[str(year)] = []
        start_month = 12 if year != current_year else current_month
        for month in range(start_month, 0, -1):
            month_str = f"{year}-{month:02d}"
            month_name_ru = get_russian_month_name(month)
            available_months[str(year)].append({"value": month_str, "name": month_name_ru})

    return available_months


def get_russian_month_name(month):
    month_names_ru = {
        1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель",
        5: "Май", 6: "Июнь", 7: "Июль", 8: "Август",
        9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь"
    }
    return month_names_ru[month]


# главная страница при входе социального работника
@app.route('/select_month')
@nocache
@login_required
@role_required('social_worker')
def select_month():
    start_year = current_user.start_date.year if current_user.start_date else None
    available_months = generate_available_months(start_year)
    return render_template('select_month.html', available_months=available_months)

# выбран месяц
@app.route('/month/<selected_month>')
@nocache
@login_required
@role_required('social_worker')
def month_view(selected_month):
    year, month = selected_month.split('-')
    num_days = calendar.monthrange(int(year), int(month))[1]
    days = [f"{year}-{month}-{d:02d}" for d in range(1, num_days + 1)]
    current_month = datetime.now().strftime("%Y-%m")
    can_edit = selected_month == current_month
    month_name_ru = get_russian_month_name(int(month))
    formatted_month = f"{month_name_ru} {year}"

    return render_template('month_view.html', selected_month=selected_month, formatted_month=formatted_month, days=days, can_edit=can_edit)



# выбран день в месяце
@app.route('/day/<selected_date>', methods=['GET', 'POST'])
@nocache
@login_required
@role_required('social_worker')
def day_view(selected_date):
    selected_date_dt = datetime.strptime(selected_date, '%Y-%m-%d')
    can_edit = (
    selected_date_dt.strftime("%Y-%m") == datetime.now().strftime("%Y-%m") and
    selected_date_dt.day <= datetime.now().day and
    selected_date_dt.weekday() != 6
    )
    daily_services = orm.get_data(Service).filter_by(date=selected_date_dt).all()
    selected_client = None

    # здесь получаем все связи между функциональными классами и типами услуг
    func_service_links = orm.get_data(FunctionalClassServiceType).all()
    all_services_by_func_class = {}

    for link in func_service_links:
        if link.functional_class_id not in all_services_by_func_class:
            all_services_by_func_class[link.functional_class_id] = []
        service_type = orm.get_by_id(ServiceType, link.service_type_id)
        if service_type:
            all_services_by_func_class[link.functional_class_id].append(service_type)

    if request.method == 'POST':
        if 'extra_client' in request.form:
            client_id = int(request.form['extra_client'])
            selected_client = orm.get_by_id(Client, client_id)
            extra_service_type_ids = request.form.getlist('extra_service_types')

            if selected_client:
                start_of_week = selected_date_dt - timedelta(days=selected_date_dt.weekday())
                end_of_week = start_of_week + timedelta(days=6)

                weekly_visits = orm.get_data(Service).filter(
                    Service.client_id == client_id,
                    Service.date >= start_of_week,
                    Service.date <= end_of_week
                ).count()

                if weekly_visits + len(extra_service_type_ids) > selected_client.visit_count:
                    flash('Превышено количество посещений в неделю для данного гражданина.', 'error')
                    return redirect(url_for('day_view', selected_date=selected_date))

                for service_type_id in extra_service_type_ids:
                    new_service = Service(
                        date=selected_date_dt,
                        service_type_id=int(service_type_id),
                        social_worker_id=current_user.id,
                        client_id=client_id
                    )
                    orm.add_inst(new_service)

                flash('Услуги успешно добавлены для дополнительного гражданина.', 'success')
                return redirect(url_for('day_view', selected_date=selected_date))

        elif 'client_id' in request.form:
            client_id = int(request.form['client_id'])
            selected_client = orm.get_by_id(Client, client_id)

            start_of_week = selected_date_dt - timedelta(days=selected_date_dt.weekday())
            end_of_week = start_of_week + timedelta(days=6)

            weekly_visits = orm.get_data(Service).filter(
                Service.client_id == client_id,
                Service.date >= start_of_week,
                Service.date <= end_of_week
            ).count()

            service_type_ids = request.form.getlist('service_types')

            if weekly_visits + len(service_type_ids) > selected_client.visit_count:
                flash('Превышено количество посещений в неделю для данного гражданина.', 'error')
                return redirect(url_for('day_view', selected_date=selected_date))

            orm.get_data(Service).filter_by(
                client_id=client_id,
                social_worker_id=current_user.id,
                date=selected_date_dt
            ).delete(synchronize_session=False)

            for service_type_id in service_type_ids:
                new_service = Service(
                    date=selected_date_dt,
                    service_type_id=int(service_type_id),
                    social_worker_id=current_user.id,
                    client_id=client_id
                )
                orm.add_inst(new_service)

            flash('Услуги успешно сохранены!', 'success')
            return redirect(url_for('day_view', selected_date=selected_date))

    assigned_clients = orm.get_data(Client).filter_by(user_id=current_user.id, status='Активен').all()
    clients_with_services = orm.get_data(Client).join(Service).filter(
        Service.social_worker_id == current_user.id,
        Service.date == selected_date_dt
    ).all()

    clients_dict = {client.id: client for client in assigned_clients}
    for client in clients_with_services:
        clients_dict.setdefault(client.id, client)

    all_clients_for_day = list(clients_dict.values())

    services_by_client = {}
    for client in all_clients_for_day:
        client_services = orm.get_data(Service).filter_by(
            client_id=client.id,
            social_worker_id=current_user.id,
            date=selected_date_dt
        ).all()
        services_by_client[client.id] = [s.service_type_id for s in client_services]

    # формируется словарь допустимых типов услуг по каждому гражданину
    client_service_types = {}
    for client in all_clients_for_day:
        func_class_id = client.func_class
        allowed_services = all_services_by_func_class.get(func_class_id, [])
        client_service_types[client.id] = [{'id': s.id, 'name': s.service_name} for s in allowed_services]


    # передача словаря в шаблон страницы:
    return render_template(
        'day_view.html',
        selected_date=selected_date,
        services=daily_services,
        can_edit=can_edit,
        assigned_clients=all_clients_for_day,
        unassigned_clients=orm.get_data(Client).filter(Client.user_id != current_user.id, Client.status=='Активен').all(),
        selected_client=selected_client,
        services_by_client=services_by_client,
        client_service_types=client_service_types
    )

# для формирования списка доступных услуг динаминески
@app.route('/get_services_for_client/<int:client_id>')
@nocache
@login_required
@role_required('social_worker')
def get_services_for_client(client_id):
    client = orm.get_by_id(Client, client_id)
    if not client:
        return jsonify({'services': []})

    func_service_links = orm.get_data(FunctionalClassServiceType).all()
    all_services_by_func_class = {}
    for link in func_service_links:
        if link.functional_class_id not in all_services_by_func_class:
            all_services_by_func_class[link.functional_class_id] = []
        service_type = orm.get_by_id(ServiceType, link.service_type_id)
        if service_type:
            all_services_by_func_class[link.functional_class_id].append(service_type)

    allowed_services = all_services_by_func_class.get(client.func_class, [])
    services_list = [{'id': s.id, 'name': s.service_name} for s in allowed_services]

    return jsonify({'services': services_list})


# формирование Дневника-отчета
@app.route('/report', methods=['POST'])
@nocache
@login_required
@role_required('social_worker')
def report():
    selected_month = request.form.get('selected_month')
    if not selected_month:
        return "Ошибка: Не указан месяц для отчета", 400

    try:
        year, month = map(int, selected_month.split('-'))
    except ValueError:
        return "Ошибка: Неверный формат месяца", 400

    document = Document(r'шаблон дневника отчета.docx') # указать документ
    month_name_ru = get_russian_month_name(month)
    header_text = f"ДНЕВНИК-ОТЧЕТ\nучета работы социального работника {current_user.surname} {current_user.first_name} {current_user.patronymic} \nза {month_name_ru}  {year} г."
    document.paragraphs[0].insert_paragraph_before(header_text, style='Normal')


    last_day = calendar.monthrange(year, month)[1]
    filtered_services = orm.get_data(Service).filter(
        Service.date.between(f'{year}-{month:02d}-01', f'{year}-{month:02d}-{last_day}'), Service.social_worker_id == current_user.id).all()
    table = document.tables[0]


    for row in table.rows[5:]:
        table.rows._tbl.remove(row._element)

    unique_client_ids = list(set(service.client_id for service in filtered_services))
    client_names, client_func_classes, client_visits = [], [], {}
    total_visits_count = 0

    for client_id in unique_client_ids:
        client = orm.get_by_id(Client, client_id)
        if not client:
            client_names.append("Неизвестный клиент")
            client_func_classes.append("Неизвестный класс")
        else:
            client_names.append(f"{client.first_name} {client.last_name}")
            client_func_classes.append(
                client.functional_class.class_name if client.functional_class else "Неизвестный класс")
        client_visits[client_id] = sum(1 for service in filtered_services if service.client_id == client_id)
        total_visits_count += client_visits[client_id]

    for i, name in enumerate(client_names):
        if i + 1 < len(table.rows[0].cells):
            table.rows[0].cells[i + 1].text = name
            table.rows[2].cells[i + 1].text = client_func_classes[i]
            table.rows[3].cells[i + 1].text = str(client_visits[unique_client_ids[i]])

    table.rows[3].cells[-1].text = str(total_visits_count)

    unique_service_names = list(set(service.service_type for service in filtered_services))
    for service_type in unique_service_names:
        row_cells = table.add_row().cells
        row_cells[0].text = service_type.service_name
        total_service_count = 0
        for j, client_id in enumerate(unique_client_ids):
            dates = [s.date.strftime('%d') for s in filtered_services if
                     s.service_type == service_type and s.client_id == client_id]
            row_cells[j + 1].text = ', '.join(dates)
            total_service_count += len(dates)
        row_cells[-1].text = str(total_service_count)

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name=f'report_{selected_month}_full.docx')


# для добавления услуг
@app.route('/add_service/<int:client_id>', methods=['POST'])
@nocache
@login_required
@role_required('social_worker')
def add_service(client_id):
    service_types = request.form.getlist('service_types')
    selected_date = request.form.get('selected_date', datetime.now().date())

    for service_type in service_types:
        if service_type:
            new_service = Service(
                date=selected_date,
                service_type_id=int(service_type),
                social_worker_id=current_user.id,
                client_id=client_id
            )
            orm.add_inst(new_service)

    return redirect(url_for('day_view', selected_date=selected_date))

# профиль соц работника
@app.route('/social')
@nocache
@login_required
@role_required('social_worker')
def profilesoc():

    user = orm.get_by_id(User, current_user.id)
    center_name = user.territorial_center.center_name if user.territorial_center else 'Не указано'
    department_name = user.department.department_name if user.department else 'Не указано'

    return render_template('profilesoc.html',
                           user=user,
                           center_name=center_name,
                           department_name=department_name)

# инспектор

# главная страница инспектора

@app.route('/inspector_welcome', methods=['GET', 'POST'])
@nocache
@login_required
@role_required('inspector')
def inspector_welcome():
    social_workers = orm.get_data(User).filter_by(role_id=2).all()
    available_months = {
        worker.id: generate_available_months(worker.start_date.year if worker.start_date else 2023)
        for worker in social_workers
    }

    months = {
        1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель",
        5: "Май", 6: "Июнь", 7: "Июль", 8: "Август",
        9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь"
    }

    return render_template('inspector_panel.html', social_workers=social_workers, available_months=available_months, months=months)

# вкладка граждане
@app.route('/citizens', methods=['GET', 'POST'])
@nocache
@login_required
@role_required('inspector')
def citizens():
    if request.method == 'POST':
        edit_id = request.form.get('edit_id')
        last_name = request.form['last_name']
        first_name = request.form['first_name']
        patronymic = request.form['patronymic']
        address = request.form['address']
        func_class = request.form['func_class']
        user_id = request.form['social_worker']
        contract_number = request.form['contract_number']

        # Даты
        def parse_date(s): return datetime.strptime(s, '%Y-%m-%d') if s else None
        contract_date = parse_date(request.form.get('contract_date'))
        birth_date = parse_date(request.form.get('birth_date'))
        contract_end_date = parse_date(request.form.get('contract_end_date'))

        # Другие поля
        visit_count_raw = request.form.get('visit_count', '0')
        visit_count = int(visit_count_raw) if visit_count_raw.isdigit() else 0
        status = request.form.get('status') or None
        gender = request.form.get('gender') or None
        phone_number = request.form.get('phone_number') or None

        if edit_id:
            citizen = orm.get_by_id(Client, edit_id)
            citizen.last_name = last_name
            citizen.first_name = first_name
            citizen.patronymic = patronymic
            citizen.address = address
            citizen.func_class = func_class
            citizen.user_id = user_id
            citizen.contract_number = contract_number
            citizen.contract_date = contract_date
            citizen.birth_date = birth_date
            citizen.contract_end_date = contract_end_date
            citizen.visit_count = visit_count
            citizen.status = status
            citizen.gender = gender
            citizen.phone_number  = phone_number
            orm.commit()
        else:
            new_citizen = Client(
                last_name=last_name,
                first_name=first_name,
                patronymic=patronymic,
                address=address,
                func_class=func_class,
                user_id=user_id,
                contract_number=contract_number,
                contract_date=contract_date,
                birth_date=birth_date,
                contract_end_date=contract_end_date,
                visit_count=visit_count,
                status=status,
                gender=gender,
                phone_number=phone_number
            )
            orm.add_inst(new_citizen)

        return redirect(url_for('citizens'))

    query = request.args.get('query', '').strip().lower()
    func_class = request.args.get('func_class')
    social_worker = request.args.get('social_worker')
    status = request.args.get('status')

    citizens_query = orm.get_data(Client)

    if query:
        citizens_query = citizens_query.filter(
            (Client.last_name.ilike(f"%{query}%")) |
            (Client.first_name.ilike(f"%{query}%")) |
            (Client.patronymic.ilike(f"%{query}%")) |
            (Client.address.ilike(f"%{query}%"))
        )
    if func_class:
        citizens_query = citizens_query.filter(Client.func_class == func_class)
    if social_worker:
        citizens_query = citizens_query.filter(Client.user_id == social_worker)
    if status:
        citizens_query = citizens_query.filter(Client.status == status)

    citizens = citizens_query.all()

    # --- Справочники ---
    functional_classes = orm.get_data(FunctionalClass).all()
    social_workers = orm.get_data(User).filter(User.role.has(role_name='social_worker')).all()
    func_map = {fc.id: fc.class_name for fc in functional_classes}
    sw_map = {sw.id: f"{sw.surname} {sw.first_name}" for sw in social_workers}

    return render_template('citizens.html',
                           citizens=citizens,
                           functional_classes=functional_classes,
                           social_workers=social_workers,
                           functional_class_map=func_map,
                           social_worker_map=sw_map)



# профиль инспектора
@app.route('/profile')
@nocache
@login_required
@role_required('inspector')
def profile():


    user = orm.get_by_id(User, current_user.id)
    center_name = user.territorial_center.center_name if user.territorial_center else 'Не указано'
    department_name = user.department.department_name if user.department else 'Не указано'

    return render_template('profile.html',
                           user=user,
                           center_name=center_name,
                           department_name=department_name)



# формирования дневника-отчета

@app.route('/inspector_generate_report', methods=['GET', 'POST'])
@nocache
@login_required
@role_required('inspector')
def inspector_generate_report():

    if request.method == 'GET':
        social_workers = orm.get_data(User).filter_by(role_id=2).all()

        # Словарь русских названий месяцев
        months = {
            1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель",
            5: "Май", 6: "Июнь", 7: "Июль", 8: "Август",
            9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь"
        }

        return render_template(
            'inspector_panel.html',
            social_workers=social_workers,
            months=months
        )

    # POST
    worker_id = request.form.get('social_worker_id')
    year = request.form.get('year')
    month = request.form.get('month')

    if not (worker_id and year and month):
        return "Ошибка: все поля обязательны", 400

    try:
        year = int(year)
        month = int(month)
    except ValueError:
        return "Ошибка: неверный формат даты", 400

    selected_worker = orm.get_by_id(User, worker_id)
    if not selected_worker:
        return "Ошибка: социальный работник не найден", 404

    document = Document(r'шаблон дневника отчета.docx') # указать документ
    month_name_ru = calendar.month_name[month]

    header_text = f"ДНЕВНИК-ОТЧЕТ\nучета работы социального работника {selected_worker.surname} {selected_worker.first_name} {selected_worker.patronymic} \nза {month_name_ru} {year} г."
    document.paragraphs[0].insert_paragraph_before(header_text, style='Normal')

    def replace_text(paragraph, key, value):
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, value)

    for paragraph in document.paragraphs:
        replace_text(paragraph, 'SOCIAL_WORKER_NAME', f"{selected_worker.surname} {selected_worker.first_name} {selected_worker.patronymic}")
        replace_text(paragraph, 'месяц отчета', month_name_ru)
        replace_text(paragraph, 'год отчета', str(year))

    last_day = calendar.monthrange(year, month)[1]

    services = orm.get_data(Service).filter(
        Service.date.between(f"{year}-{month:02d}-01", f'{year}-{month:02d}-{last_day}'),
        Service.social_worker_id == selected_worker.id
    ).all()

    table = document.tables[0]
    for row in table.rows[5:]:
        table.rows._tbl.remove(row._element)

    client_ids = list(set(s.client_id for s in services))
    client_names, func_classes, visits = [], [], {}

    total_visits_count=0

    for client_id in client_ids:
        client = orm.get_by_id(Client, client_id)
        if client:
            client_names.append(f"{client.first_name} {client.last_name}")
            func_classes.append(client.functional_class.class_name if client.functional_class else "Неизвестно")
        else:
            client_names.append("Неизвестно")
            func_classes.append("Неизвестно")
        visits[client_id] = sum(1 for s in services if s.client_id == client_id)
        total_visits_count += visits[client_id]

    table.rows[3].cells[-1].text = str(total_visits_count)

    for i, name in enumerate(client_names):
        if i + 1 < len(table.rows[0].cells):
            table.rows[0].cells[i + 1].text = name
            table.rows[2].cells[i + 1].text = func_classes[i]
            table.rows[3].cells[i + 1].text = str(visits[client_ids[i]])

    unique_service_types = list(set(s.service_type for s in services))
    for s_type in unique_service_types:
        row = table.add_row().cells
        row[0].text = s_type.service_name
        total = 0
        for j, client_id in enumerate(client_ids):
            days = [s.date.strftime('%d') for s in services if s.client_id == client_id and s.service_type == s_type]
            row[j + 1].text = ', '.join(days)
            total += len(days)
        row[-1].text = str(total)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(buffer,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True,
                     download_name=f'отчет_{month_name_ru}_{year}_{selected_worker.surname}.docx')

# вкладка Акты и формирование актов
@app.route('/acts', methods=['GET', 'POST'])
@nocache
@login_required
@role_required('inspector')
def acts_view():

    citizens = orm.get_data(Client).all()

    if request.method == 'POST':
        year = request.form.get('year')
        month = request.form.get('month')
        citizen_id = request.form.get('citizen')

        if not year or not month or not citizen_id:
            return "Пожалуйста, заполните все поля", 400

        citizen = orm.get_by_id(Client, citizen_id)
        if not citizen:
            return "Гражданин не найден", 404

        center = orm.get_by_id(TerritorialCenter, citizen.user.territorial_center_id)

        document = Document((r'шаблонакт.docx')) # указать документ

        replace_placeholders(document, citizen, center, year, month)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return send_file(buffer,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True,
                         download_name=f'act_{citizen.first_name}_{citizen.last_name}_{year}_{month}.docx')

    return render_template('acts.html', citizens=citizens)


def replace_placeholders(document, citizen, center, year, month):
    # Get the current date for the act
    dateact = datetime.now().strftime('%d-%m-%Y')
    order_number = citizen.contract_number
    contract_date = citizen.contract_date.strftime('%d-%m-%Y') if citizen.contract_date else 'Не указано'

    months_mapping = {
        'Январь': 1,
        'Февраль': 2,
        'Март': 3,
        'Апрель': 4,
        'Май': 5,
        'Июнь': 6,
        'Июль': 7,
        'Август': 8,
        'Сентябрь': 9,
        'Октябрь': 10,
        'Ноябрь': 11,
        'Декабрь': 12
    }
    month_name = request.form.get('month')
    month = months_mapping.get(month_name)
    if month is None:
        return "Ошибка: Неверное название месяца", 400
    year = int(request.form.get('year'))

    first_day = datetime(year, month, 1)
    if month == 12:
        last_day = datetime(year + 1, 1, 1)
    else:
        last_day = datetime(year, month + 1, 1)

    services = orm.get_data(Service).filter(
        Service.client_id == citizen.id,
        Service.date >= first_day,
        Service.date < last_day
    ).all()

    service_names = [s.service_type.service_name for s in services]
    service_dates = sorted(set([s.date.strftime('%d') for s in services]))


    social_worker_ids = set(s.social_worker_id for s in services if s.social_worker_id)
    social_workers = [orm.get_by_id(User, sw_id) for sw_id in social_worker_ids]
    social_worker_names = []
    for sw in social_workers:
        if sw:
            initials = f"{sw.first_name[0]}." if sw.first_name else ""
            patronymic_initial = f"{sw.patronymic[0]}." if sw.patronymic else ""
            social_worker_names.append(f"{sw.surname} {initials}{patronymic_initial}")

    social_worker_name_str = ", ".join(social_worker_names) if social_worker_names else "Не указано"

    for paragraph in document.paragraphs:
        paragraph.text = paragraph.text.replace('{{center_name}}', center.center_name)
        paragraph.text = paragraph.text.replace('{{director_name}}', center.director_full_name)
        paragraph.text = paragraph.text.replace('{{dateact}}', dateact)
        paragraph.text = paragraph.text.replace('{{citizen_name}}', f" {citizen.last_name} {citizen.first_name} {citizen.patronymic}")
        paragraph.text = paragraph.text.replace('{{address}}', citizen.address)
        paragraph.text = paragraph.text.replace('{{order_number}}', order_number)
        paragraph.text = paragraph.text.replace('{{date}}', contract_date)
        paragraph.text = paragraph.text.replace('{{services}}', ', '.join(service_names) if service_names else 'Нет услуг')
        paragraph.text = paragraph.text.replace('{{dates_of_service}}', ', '.join(service_dates) if service_dates else 'Нет дат')
        paragraph.text = paragraph.text.replace('{{social_worker}}', social_worker_name_str)

        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

        if '{{services}}' in paragraph.text or '{{dates_of_service}}' in paragraph.text:
            for run in paragraph.runs:
                run.font.underline = True

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

    return document
if __name__ == '__main__':
    app.run(debug=True)
